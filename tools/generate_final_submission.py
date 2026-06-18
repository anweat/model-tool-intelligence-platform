from __future__ import annotations

import json
import math
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION = ROOT / "docs" / "submission" / "final"
IMAGE_DIR = SUBMISSION / "images"
MODULE_DIR = IMAGE_DIR / "modules"
PACKAGE_ROOT = SUBMISSION / "package"
PACKAGE_NAME = "学号_姓名_模型工具调用智能分析平台"
PACKAGE_DIR = PACKAGE_ROOT / PACKAGE_NAME
TEACHER_LINK = (
    "https://devcloud.cn-north-4.huaweicloud.com/DevUCUsermgmtPortal/applyjoinmember/"
    "e590e9bcd23e4b99b9d9052a0b3fa983/join?"
    "code=QUFBQUFnQUFBQUFBQUFBQUFBQUFBUUFBQUFyZGRuR0N0TDV6bkpUQXNkV3pOaVlQZklEeGx6c2d6VVRlTUVHS0JtYWUrUVFBQUFFQUFBQUFBQUFBWU1DM1RkempXbHUrdktSVk92V2Y2WEl1STBLK0pEWStTdUF0UEZSMGlmcEdpYk92Z1dOd1RjdFNsb0JZOVF4RzZFaGdtOHFLenpOTWdWZEd2NXltbnZOaGptS1RvSnVmdERyWWd0dTZpRUl3Y1ZhdTlaSUdjdHdKRzFpdnMwbVJnQT09"
)

FONT_PATH = Path("C:/Windows/Fonts/msyh.ttc")
FONT_BOLD_PATH = Path("C:/Windows/Fonts/msyhbd.ttc")


MODULES = [
    {
        "key": "provider",
        "name": "供应商管理模块",
        "goal": "维护模型供应商基础资料、接口地址、启停状态和关联模型数量，为 API Key、模型目录和调用记录提供上游主数据。",
        "classes": [
            ("Provider", ["id", "name", "baseUrl", "status", "notes"]),
            ("StoreRepository", ["loadStore()", "saveStore()", "appendAudit()"]),
            ("AuditLog", ["id", "action", "targetType", "targetId", "createdAt"]),
        ],
        "flow": ["填写供应商名称", "校验名称与 Base URL", "写入 providers", "记录审计日志", "刷新供应商列表"],
        "sequence": ["前端提交供应商表单", "API 校验必填字段", "Store 写入 Provider", "AuditLog 记录 create_provider", "返回 state 与 summary"],
    },
    {
        "key": "apikey",
        "name": "API Key 管理模块",
        "goal": "按供应商保存调用凭证，前端与报表只展示脱敏密钥，降低敏感信息泄露风险。",
        "classes": [
            ("ApiKey", ["id", "providerId", "label", "owner", "maskedSecret", "status"]),
            ("Provider", ["id", "name", "status"]),
            ("SecurityPolicy", ["maskSecret()", "stripSecretFromExport()"]),
        ],
        "flow": ["选择供应商", "录入 Key 标签和负责人", "保存完整 secret", "接口返回 maskedSecret", "导出报表移除 secret"],
        "sequence": ["前端提交 Key", "API 校验供应商存在", "Store 保存 secret", "publicStore 生成 maskedSecret", "前端展示脱敏结果"],
    },
    {
        "key": "model",
        "name": "模型目录与价格模块",
        "goal": "维护模型上下文窗口、输入/输出单价、缓存折扣和工具调用能力，为计费与可行性测算提供基础参数。",
        "classes": [
            ("Model", ["id", "providerId", "name", "contextWindow", "inputPricePer1K", "outputPricePer1K", "cacheDiscount"]),
            ("Provider", ["id", "name"]),
            ("PricingRule", ["validatePrice()", "validateDiscount()", "calculateUnitCost()"]),
        ],
        "flow": ["选择供应商", "填写模型能力", "校验价格与折扣范围", "保存模型目录", "参与调用计费"],
        "sequence": ["前端新增模型", "API 校验 providerId", "API 校验数值范围", "Store 写入 Model", "刷新模型下拉与价格表"],
    },
    {
        "key": "call",
        "name": "调用记录与缓存模块",
        "goal": "登记每次模型工具调用的 Token、耗时、缓存命中、成功状态和失败类型，形成后续统计分析的数据基础。",
        "classes": [
            ("CallRecord", ["id", "providerId", "modelId", "toolName", "inputTokens", "outputTokens", "cacheHit", "latencyMs"]),
            ("Model", ["id", "providerId", "cacheDiscount"]),
            ("CallValidator", ["validateProviderModel()", "validateToken()", "enrichCall()"]),
        ],
        "flow": ["选择供应商和模型", "填写工具与 Token", "校验模型归属", "计算费用和缓存节省", "写入调用记录"],
        "sequence": ["前端提交调用", "API 校验 provider/model 关系", "API 读取模型价格", "enrichCall 计算成本", "返回最新总览"],
    },
    {
        "key": "billing",
        "name": "统计报表与计费模块",
        "goal": "聚合调用次数、Token、缓存命中率、成功率、供应商/模型费用、工具摘要和预算告警，支撑课程报告和审计导出。",
        "classes": [
            ("SummaryService", ["computeSummary()", "groupByProvider()", "groupByTool()", "budgetStatus()"]),
            ("BillingPolicy", ["monthlyBudget", "alertThreshold", "retentionDays"]),
            ("ReportExporter", ["buildExport()", "stripSecrets()", "exportedAt"]),
        ],
        "flow": ["读取调用记录", "按模型价格计算总费用", "汇总缓存与成功率", "判断预算告警", "生成 JSON 报表"],
        "sequence": ["前端请求 /api/state", "API 读取 store", "SummaryService 聚合指标", "BudgetPolicy 生成状态", "前端刷新仪表盘"],
    },
    {
        "key": "insight",
        "name": "智能评估与治理模块",
        "goal": "把工具调用数据转化为系统画像、工具价值等级、费用预测、上线可行性评分和治理建议，体现项目创新点。",
        "classes": [
            ("InsightEngine", ["computeInsights()", "feasibilityScore()", "toolValueLevel()", "governanceAdvice()"]),
            ("ToolInsight", ["toolName", "calls", "successRate", "cacheHitRate", "averageLatencyMs", "valueLevel"]),
            ("FeasibilityRequest", ["modelId", "monthlyCalls", "averageInputTokens", "averageOutputTokens", "targetCacheHitRate"]),
        ],
        "flow": ["读取统计摘要", "识别主场景和主模型", "计算工具价值等级", "评估预算/延迟/成功率", "输出上线建议"],
        "sequence": ["前端请求 /api/insights", "InsightEngine 聚合工具", "计算可行性评分", "生成治理建议", "返回系统画像"],
    },
]


def ensure_dirs() -> None:
    for path in [SUBMISSION, IMAGE_DIR, MODULE_DIR, PACKAGE_ROOT]:
        path.mkdir(parents=True, exist_ok=True)


def generate_mvp_segments() -> list[Path]:
    source = ROOT / "docs" / "submission" / "images" / "mvp-dashboard.png"
    if not source.exists():
        return []
    image = Image.open(source).convert("RGB")
    width, height = image.size
    cuts = [
        ("mvp-dashboard-01-overview.png", "总览、预算与统计区", 0, min(1050, height)),
        ("mvp-dashboard-02-management.png", "工具洞察、供应商和 API Key 管理区", min(900, height), min(2050, height)),
        ("mvp-dashboard-03-model-call.png", "模型目录与调用记录区", min(1900, height), height),
    ]
    outputs: list[Path] = []
    for file_name, label, top, bottom in cuts:
        if bottom <= top:
            continue
        crop = image.crop((0, top, width, bottom))
        canvas = Image.new("RGB", (width, crop.height + 88), "#f8fafc")
        draw = ImageDraw.Draw(canvas)
        draw.text((42, 28), f"MVP 页面分段截图：{label}", font=font(30, True), fill="#0f172a")
        canvas.paste(crop, (0, 88))
        out = IMAGE_DIR / file_name
        canvas.save(out)
        outputs.append(out)
    return outputs


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_BOLD_PATH if bold and FONT_BOLD_PATH.exists() else FONT_PATH
    return ImageFont.truetype(str(path), size)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, ft: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        trial = current + ch
        if draw.textbbox((0, 0), trial, font=ft)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, ft: ImageFont.FreeTypeFont, fill=(20, 35, 50)) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, ft, max(40, x2 - x1 - 20))
    line_h = ft.size + 8
    y = y1 + ((y2 - y1) - len(lines) * line_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=ft)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, font=ft, fill=fill)
        y += line_h


def rounded_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], label: str, fill: str, outline: str, ft=None) -> None:
    draw.rounded_rectangle(xy, radius=12, fill=fill, outline=outline, width=3)
    draw_centered_text(draw, xy, label, ft or font(24, True))


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill="#334155") -> None:
    draw.line([start, end], fill=fill, width=3)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 12
    p1 = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
    p2 = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
    draw.polygon([end, p1, p2], fill=fill)


def title(draw: ImageDraw.ImageDraw, text: str, width: int) -> None:
    ft = font(34, True)
    bbox = draw.textbbox((0, 0), text, font=ft)
    draw.text(((width - (bbox[2] - bbox[0])) / 2, 28), text, font=ft, fill="#0f172a")


def save_flow(module: dict) -> Path:
    width, height = 1400, 520
    img = Image.new("RGB", (width, height), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title(draw, f"{module['name']}流程图", width)
    steps = module["flow"]
    box_w, box_h, gap = 210, 92, 34
    start_x = (width - (len(steps) * box_w + (len(steps) - 1) * gap)) // 2
    y = 210
    for idx, step in enumerate(steps):
        x = start_x + idx * (box_w + gap)
        rounded_box(draw, (x, y, x + box_w, y + box_h), step, "#ecfdf5", "#059669", font(23, True))
        if idx < len(steps) - 1:
            arrow(draw, (x + box_w, y + box_h // 2), (x + box_w + gap - 8, y + box_h // 2))
    out = MODULE_DIR / f"{module['key']}-flow.png"
    img.save(out)
    return out


def save_class(module: dict) -> Path:
    width, height = 1400, 680
    img = Image.new("RGB", (width, height), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title(draw, f"{module['name']}类图", width)
    positions = [(90, 170), (515, 170), (940, 170)]
    class_w = 330
    for (name, fields), (x, y) in zip(module["classes"], positions):
        header_h = 58
        body_h = 270
        draw.rounded_rectangle((x, y, x + class_w, y + header_h + body_h), radius=12, fill="#ffffff", outline="#2563eb", width=3)
        draw.rounded_rectangle((x, y, x + class_w, y + header_h), radius=12, fill="#dbeafe", outline="#2563eb", width=3)
        draw_centered_text(draw, (x, y, x + class_w, y + header_h), name, font(24, True))
        fy = y + header_h + 26
        for field in fields:
            draw.text((x + 24, fy), f"+ {field}", font=font(21), fill="#0f172a")
            fy += 42
    arrow(draw, (420, 330), (515, 330))
    arrow(draw, (845, 330), (940, 330))
    draw.text((555, 455), "关联/调用", font=font(21), fill="#475569")
    out = MODULE_DIR / f"{module['key']}-class.png"
    img.save(out)
    return out


def save_sequence(module: dict) -> Path:
    width, height = 1400, 720
    img = Image.new("RGB", (width, height), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title(draw, f"{module['name']}时序图", width)
    lanes = [("前端页面", 180), ("API 服务", 460), ("JSON 存储", 740), ("领域服务", 1020), ("页面状态", 1240)]
    for name, x in lanes:
        rounded_box(draw, (x - 85, 125, x + 85, 175), name, "#eef2ff", "#94a3b8", font(21, True))
        draw.line((x, 175, x, 650), fill="#cbd5e1", width=3)
    y = 230
    lane_x = [x for _, x in lanes]
    pairs = [(0, 1), (1, 2), (1, 3), (3, 1), (1, 4)]
    for step, (a, b) in zip(module["sequence"], pairs):
        x1, x2 = lane_x[a], lane_x[b]
        arrow(draw, (x1, y), (x2, y))
        tx = min(x1, x2) + abs(x2 - x1) / 2
        wrapped = wrap_text(draw, step, font(20), abs(x2 - x1) - 30)
        for i, line in enumerate(wrapped[:2]):
            bbox = draw.textbbox((0, 0), line, font=font(20))
            draw.text((tx - (bbox[2] - bbox[0]) / 2, y - 36 + i * 24), line, font=font(20), fill="#0f172a")
        y += 88
    out = MODULE_DIR / f"{module['key']}-sequence.png"
    img.save(out)
    return out


def generate_module_diagrams() -> dict[str, dict[str, Path]]:
    result = {}
    for module in MODULES:
        result[module["key"]] = {
            "flow": save_flow(module),
            "class": save_class(module),
            "sequence": save_sequence(module),
        }
    return result


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "AAB7C4")
        borders.append(tag)
    tblPr.append(borders)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    table.autofit = False
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        shade_cell(table.rows[0].cells[i], "E8EEF5")
        if widths:
            table.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def set_run_font(run, size=11, bold=False, color="000000") -> None:
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_p(doc: Document, text: str = "", style: str | None = None, after=6, before=0, bold=False) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        set_run_font(run, bold=bold)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_image(doc: Document, path: Path, caption: str, width_in: float = 6.2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9, color="475569")


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)
    for style_name, size, color in [
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        st = styles[style_name]
        st.font.name = "微软雅黑"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "模型工具调用智能分析平台项目报告"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if header.runs:
        set_run_font(header.runs[0], size=9, color="64748B")
    footer = section.footer.paragraphs[0]
    footer.text = "CodeArts 课程作业交付材料"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer.runs:
        set_run_font(footer.runs[0], size=9, color="64748B")


def add_cover(doc: Document) -> None:
    add_p(doc, "", after=60)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("模型工具调用智能分析平台")
    set_run_font(r, size=26, bold=True, color="0B2545")
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目报告与课程作业交付说明")
    set_run_font(r, size=15, color="475569")
    p.paragraph_format.space_after = Pt(42)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["项目名称", "模型工具调用智能分析平台"],
            ["提交人", "学号/姓名：提交前替换"],
            ["平台", "华为云 CodeArts"],
            ["报告版本", "V1.0"],
            ["生成日期", "2026-06-18"],
            ["教师浏览链接", TEACHER_LINK],
        ],
        [1.6, 4.7],
    )
    add_p(
        doc,
        "说明：本文档依据课程作业要求编写，覆盖需求管理、软件建模、概要设计、详细设计、编码实现、测试验收、AI 辅助说明和提交材料清单。",
        after=8,
    )
    doc.add_page_break()


def build_doc(diagrams: dict[str, dict[str, Path]]) -> Path:
    doc = Document()
    configure_doc(doc)
    add_cover(doc)

    doc.add_heading("1 项目概述", level=1)
    add_p(doc, "本项目面向多模型应用开发场景，统一管理模型供应商、API Key、模型目录与调用日志，并基于工具调用数据自动生成系统画像、工具价值分析、成本风险、缓存收益和上线可行性评估。")
    add_bullets(
        doc,
        [
            "业务目标：将分散的模型调用凭证、模型能力和调用成本统一纳入可审计的管理流程。",
            "技术目标：使用轻量 Node.js HTTP 服务、本地 JSON 持久化和浏览器前端完成课程 MVP。",
            "创新点：把工具调用日志转化为主场景、主模型、工具价值等级、预算风险和上线可行性建议。",
            "交付目标：满足 CodeArts 需求管理、软件建模、代码仓库、实现报告和最终压缩包提交要求。",
        ],
    )

    doc.add_heading("2 需求管理与需求分析", level=1)
    add_p(doc, "项目在 CodeArts Req 中按 Epic、Feature、User Story 组织需求，共形成 1 个 Epic、6 个 Feature 和 18 条 User Story。工作项已闭环至已关闭状态，满足需求管理可追踪要求。")
    add_table(
        doc,
        ["层级", "数量", "说明"],
        [
            ["Epic", "1", "统一模型供应商 API Key 与调用成本管理平台"],
            ["Feature", "6", "供应商、API Key、模型目录、调用记录、统计报表、智能评估"],
            ["User Story", "18", "每条 Story 包含角色、目标、收益、优先级和验收标准"],
        ],
        [1.2, 1.1, 4.0],
    )
    req_img = ROOT / "docs" / "submission" / "images" / "requirements-status.png"
    if req_img.exists():
        add_image(doc, req_img, "图 1 CodeArts 需求管理状态截图", 6.2)
    add_table(
        doc,
        ["Feature", "对应用户故事", "核心验收点"],
        [
            ["供应商管理", "US-01 ~ US-02", "供应商新增、启停、模型数量统计"],
            ["API Key 管理", "US-03 ~ US-05", "Key 保存、脱敏展示、状态管理"],
            ["模型目录与价格", "US-06 ~ US-08", "上下文窗口、价格、缓存折扣、工具调用能力"],
            ["调用记录与缓存", "US-09 ~ US-10", "调用登记、缓存命中、Token 和耗时采集"],
            ["统计报表与计费", "US-11 ~ US-14", "总览统计、预算告警、JSON 报表导出"],
            ["智能洞察与可行性", "US-15 ~ US-18", "系统画像、工具价值、费用预测、治理建议"],
        ],
        [1.55, 1.65, 3.1],
    )
    doc.add_heading("2.1 非功能需求", level=2)
    add_bullets(
        doc,
        [
            "安全性：接口和导出报表不得暴露完整 API Key，前端仅展示 maskedSecret。",
            "可审计性：新增、修改、删除、归档等操作写入 auditLogs。",
            "可演示性：项目无需外部数据库，默认演示数据支持离线验收。",
            "可维护性：REST API 边界清晰，统计、洞察和持久化逻辑集中在后端。",
            "可扩展性：后续可将 JSON 存储替换为 SQLite、PostgreSQL 或 MySQL，并加入账号权限。",
        ],
    )

    doc.add_heading("3 软件建模", level=1)
    add_p(doc, "软件建模交付包括用例图、活动图、类图、顺序图和模块级补充图。正式报告插图统一采用 PNG，以避免 SVG 字体回退导致中文显示异常。")
    general_images = [
        ("use-case-key-model.png", "图 2 用例图：API Key 与模型目录管理"),
        ("use-case-usage-billing.png", "图 3 用例图：调用统计与智能分析"),
        ("activity-call-billing.png", "图 4 活动图：模型调用登记与计费"),
        ("class-diagram.png", "图 5 类图：核心领域模型"),
        ("sequence-create-call.png", "图 6 顺序图：新增调用记录"),
    ]
    for file_name, caption in general_images:
        path = ROOT / "docs" / "diagrams" / "exported-png" / file_name
        if path.exists():
            add_image(doc, path, caption, 6.2)

    doc.add_heading("4 概要设计", level=1)
    add_p(doc, "系统采用轻量单体架构，由浏览器前端、Node.js HTTP API、JSON 文件持久化和统计分析模块组成。该方案避免外部服务依赖，适合课程作业演示、CodeArts 仓库托管和本地验收。")
    add_table(
        doc,
        ["层次", "实现文件", "职责"],
        [
            ["表现层", "public/index.html, public/app.js, public/styles.css", "表单、列表、仪表盘、预算设置、方案测算和报表导出"],
            ["接口层", "server.js /api/*", "处理 CRUD、状态摘要、智能洞察、可行性测算和导出请求"],
            ["领域逻辑", "server.js 中统计与校验函数", "计算费用、缓存收益、工具洞察、系统画像和可行性评分"],
            ["持久化层", "data/store.json", "保存供应商、Key、模型、调用记录、审计日志和系统设置"],
            ["支撑材料", "docs/, tools/", "生成需求文档、建模材料、XMI、模型工程包、测试脚本和提交报告"],
        ],
        [1.0, 2.0, 3.3],
    )
    mvp_segments = generate_mvp_segments()
    if mvp_segments:
        for idx, segment in enumerate(mvp_segments, start=1):
            captions = {
                1: "图 7-1 MVP 分段截图：总览、预算与统计区",
                2: "图 7-2 MVP 分段截图：工具洞察、供应商和 API Key 管理区",
                3: "图 7-3 MVP 分段截图：模型目录与调用记录区",
            }
            add_image(doc, segment, captions[idx], 6.0)

    doc.add_heading("5 详细设计", level=1)
    add_p(doc, "详细设计按业务模块展开。每个模块均给出设计目标、核心类图、处理流程图和典型时序图，便于对应 CodeArts 软件建模和代码实现。")
    figure_no = 8
    for module in MODULES:
        doc.add_heading(f"5.{MODULES.index(module)+1} {module['name']}", level=2)
        add_p(doc, module["goal"])
        add_table(
            doc,
            ["设计项", "说明"],
            [
                ["主要输入", "前端表单、已有 store.json 数据、系统设置或调用日志"],
                ["主要输出", "最新 state、summary、insights、auditLogs 或导出报表"],
                ["关键约束", "必填字段校验、资源归属校验、数值范围校验、敏感字段脱敏"],
            ],
            [1.35, 4.95],
        )
        add_image(doc, diagrams[module["key"]]["class"], f"图 {figure_no} {module['name']}类图", 6.2)
        figure_no += 1
        add_image(doc, diagrams[module["key"]]["flow"], f"图 {figure_no} {module['name']}流程图", 6.2)
        figure_no += 1
        add_image(doc, diagrams[module["key"]]["sequence"], f"图 {figure_no} {module['name']}时序图", 6.2)
        figure_no += 1

    doc.add_heading("6 编码实现", level=1)
    add_p(doc, "系统代码已覆盖需求规划中的核心业务流程。后端使用 Node.js 内置 http 模块实现 REST API，不依赖外部数据库；前端使用原生 HTML、CSS、JavaScript 实现交互，便于在课程环境中快速运行。")
    add_table(
        doc,
        ["功能", "实现位置", "说明"],
        [
            ["供应商管理", "server.js, public/app.js", "新增、启停、删除、关联模型数量统计"],
            ["API Key 管理", "server.js", "保存 secret，接口返回 publicStore 时只保留 maskedSecret"],
            ["模型目录", "server.js", "维护价格、上下文窗口、缓存折扣和工具调用能力"],
            ["调用登记", "server.js", "校验 provider/model 归属，计算费用、缓存节省和审计日志"],
            ["统计与计费", "computeSummary()", "汇总调用次数、Token、费用、缓存命中率、成功率和预算状态"],
            ["智能洞察", "computeInsights()", "生成系统画像、工具洞察、失败类型统计和治理建议"],
            ["可行性测算", "/api/feasibility", "根据预计调用量和 Token 规模预测月费用与预算占用"],
            ["测试", "tools/smoke-test.js", "使用临时 DATA_DIR 验证关键接口和安全约束"],
        ],
        [1.4, 1.8, 3.1],
    )

    doc.add_heading("7 测试验收", level=1)
    add_p(doc, "已执行 npm test 冒烟测试，结果为 smoke test passed。测试会启动临时服务并使用临时数据目录，避免污染提交用演示数据。")
    add_table(
        doc,
        ["验证项", "结果"],
        [
            ["健康检查 /api/health", "通过"],
            ["演示数据重置 /api/reset", "通过"],
            ["供应商、Key、模型、调用记录新增", "通过"],
            ["API Key 脱敏与导出不泄露完整密钥", "通过"],
            ["供应商与模型归属校验", "通过"],
            ["可行性测算与异常输入拦截", "通过"],
            ["统计摘要、系统画像、工具洞察、报表导出", "通过"],
            ["npm test", "smoke test passed"],
        ],
        [3.2, 3.1],
    )

    doc.add_heading("8 CodeArts 交付与老师浏览链接", level=1)
    add_p(doc, "项目已在 CodeArts 中创建需求工作项，并通过 GitHub 镜像方式导入软件仓库。老师可通过以下浏览者链接申请加入项目查看需求管理、仓库和建模材料。")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("老师浏览者链接：")
    set_run_font(run, bold=True)
    add_hyperlink(p, "打开 CodeArts 项目加入链接", TEACHER_LINK)
    add_p(doc, TEACHER_LINK, after=8)
    add_table(
        doc,
        ["交付项", "位置或说明"],
        [
            ["CodeArts 项目", "项目 ID：e590e9bcd23e4b99b9d9052a0b3fa983"],
            ["GitHub 仓库", "https://github.com/anweat/model-tool-intelligence-platform"],
            ["模型工程包", "docs/codearts-import/model_tool_intelligence_master.zip"],
            ["正式建模图片", "docs/diagrams/exported-png/"],
            ["最终报告", "docs/submission/final/项目报告_模型工具调用智能分析平台.docx"],
        ],
        [1.6, 4.7],
    )

    doc.add_heading("9 生成式人工智能辅助说明", level=1)
    add_p(doc, "本作业使用生成式人工智能技术辅助完成部分工程工作，主要工具为 OpenAI Codex / ChatGPT 类代码助手。AI 辅助不替代课程设计目标判断，项目主题、功能取舍、CodeArts 导入确认、最终材料提交由本人确认。")
    add_table(
        doc,
        ["使用工具", "辅助内容", "人工确认方式"],
        [
            ["OpenAI Codex / ChatGPT 类代码助手", "辅助编写 Node.js 服务、前端交互、测试脚本、XMI/模型工程生成脚本和报告草稿", "人工阅读代码逻辑，执行 npm test，并在 CodeArts 中导入与复核"],
            ["CodeArts Agent 与华为云 SDK", "辅助创建和复核 CodeArts 工作项、仓库镜像状态", "通过 API 返回结果和页面导入结果确认"],
            ["本地脚本/Playwright/LibreOffice", "辅助生成截图、建模图片、Word 报告和渲染检查", "人工检查渲染页，确认中文、表格、图片无明显错误"],
        ],
        [1.75, 2.65, 1.9],
    )

    doc.add_heading("10 提交材料清单", level=1)
    add_table(
        doc,
        ["目录或文件", "内容说明"],
        [
            ["项目报告_模型工具调用智能分析平台.docx/.doc", "完整项目报告，含需求截图、建模图片、设计、实现、测试和 AI 辅助说明"],
            ["README.md", "项目运行说明和作业对应关系"],
            ["server.js, public/, data/", "源代码、前端页面和演示数据"],
            ["docs/codearts-requirements.md", "需求分析说明书和 User Story 清单"],
            ["docs/project-specification.md", "项目说明、接口、验收标准和创新点"],
            ["docs/implementation-report.md", "编码实现报告"],
            ["docs/diagrams/, docs/xmi/, docs/codearts-import/", "建模源文件、正式图片、XMI 和 CodeArts 模型工程包"],
            ["tools/", "测试、图形生成、XMI 生成和模型工程包生成脚本"],
            ["老师浏览链接.txt", "CodeArts 浏览者加入链接"],
        ],
        [2.25, 4.05],
    )
    add_p(doc, "最终压缩包命名为“学号_姓名_模型工具调用智能分析平台.zip”。当前生成文件使用占位前缀，提交前需替换为个人真实学号和姓名。", bold=True)

    out = SUBMISSION / "项目报告_模型工具调用智能分析平台.docx"
    doc.save(out)
    shutil.copy2(out, SUBMISSION / "report_model_tool_intelligence_final.docx")
    return out


def copy_for_package() -> None:
    if PACKAGE_DIR.exists():
        shutil.rmtree(PACKAGE_DIR)
    PACKAGE_DIR.mkdir(parents=True)
    docs_out = PACKAGE_DIR / "文档与报告"
    src_out = PACKAGE_DIR / "源代码"
    model_out = PACKAGE_DIR / "软件建模材料"
    support_out = PACKAGE_DIR / "支撑材料"
    for p in [docs_out, src_out, model_out, support_out]:
        p.mkdir(parents=True, exist_ok=True)

    for name in ["项目报告_模型工具调用智能分析平台.docx", "项目报告_模型工具调用智能分析平台.doc"]:
        src = SUBMISSION / name
        if src.exists():
            shutil.copy2(src, docs_out / name)
    for name in ["README.md"]:
        shutil.copy2(ROOT / name, src_out / name)
    for file in ["server.js", "package.json"]:
        shutil.copy2(ROOT / file, src_out / file)
    shutil.copytree(ROOT / "public", src_out / "public", dirs_exist_ok=True)
    data_out = src_out / "data"
    data_out.mkdir(parents=True, exist_ok=True)
    for file in ["store.json", "demo-call.json", "settings-patch.json"]:
        src = ROOT / "data" / file
        if src.exists():
            shutil.copy2(src, data_out / file)

    for file in ["codearts-requirements.md", "project-specification.md", "implementation-report.md", "modeling.md", "modeling-checklist.md", "codearts-operation-log.md"]:
        shutil.copy2(ROOT / "docs" / file, docs_out / file)
    report_images = docs_out / "report-images"
    report_images.mkdir(parents=True, exist_ok=True)
    for image in IMAGE_DIR.glob("*.png"):
        shutil.copy2(image, report_images / image.name)
    shutil.copytree(ROOT / "docs" / "diagrams", model_out / "diagrams", dirs_exist_ok=True)
    shutil.copytree(MODULE_DIR, model_out / "module-diagrams", dirs_exist_ok=True)
    shutil.copytree(ROOT / "docs" / "xmi", model_out / "xmi", dirs_exist_ok=True)
    shutil.copytree(ROOT / "docs" / "codearts-import", model_out / "codearts-import", dirs_exist_ok=True)
    shutil.copytree(ROOT / "tools", support_out / "tools", dirs_exist_ok=True, ignore=shutil.ignore_patterns("__pycache__", "codearts_api.py", "create_codearts_requirements.py", "upload_codearts_*.py"))
    (PACKAGE_DIR / "老师浏览链接.txt").write_text(TEACHER_LINK + "\n", encoding="utf-8")
    (PACKAGE_DIR / "提交前重命名说明.txt").write_text("请将压缩包文件名中的“学号_姓名”替换为个人真实学号和姓名。\n", encoding="utf-8")


def zip_package() -> Path:
    zip_path = SUBMISSION / f"{PACKAGE_NAME}.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in PACKAGE_DIR.rglob("*"):
            zf.write(path, path.relative_to(PACKAGE_ROOT))
    return zip_path


def main() -> None:
    ensure_dirs()
    diagrams = generate_module_diagrams()
    report = build_doc(diagrams)
    copy_for_package()
    zip_path = zip_package()
    print(json.dumps({"report": str(report), "package": str(zip_path), "teacher_link": TEACHER_LINK}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
